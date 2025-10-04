import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, urlunparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import time


class Odoo18ContributingScraper:
    def __init__(self, start_url, delay=0.8, retry=2, timeout=30):
        # Початкові параметри
        self.start_url = start_url
        parsed = urlparse(self.start_url)
        self.base_scheme = parsed.scheme or 'https'
        self.base_domain = parsed.netloc
        self.base_path_prefix = '/documentation/18.0/contributing'
        self.visited = set()
        self.ordered_urls = []
        self.doc = Document()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        self.delay = delay
        self.retry = retry
        self.timeout = timeout
        # Перевіряємо/створюємо необхідні стилі для DOCX
        self._ensure_styles()

    # =========================
    #  URL: нормалізація/перевірка
    # =========================
    def normalize_url(self, url, current_url=None):
        """
        Нормалізує URL у межах того ж домену та гілки документації:
        - абсолютний
        - без параметрів і якорів
        - тільки /documentation/18.0/contributing*
        - без статичних файлів
        """
        if not url:
            return None

        low = url.strip().lower()
        if low.startswith(('mailto:', 'javascript:', 'tel:')):
            return None

        base = current_url or self.start_url
        absolute = urljoin(base, url)
        parsed = urlparse(absolute)
        cleaned = parsed._replace(fragment='', query='')

        if cleaned.netloc != self.base_domain:
            return None

        path = cleaned.path or '/'
        if not path.startswith('/'):
            path = '/' + path

        # Ігноруємо статичні файли
        if path.endswith(('.pdf', '.zip', '.tar', '.gz', '.md')):
            return None

        # Обмеження на потрібну гілку документації
        if not path.startswith(self.base_path_prefix):
            return None

        # Уніфікація трейлінг-слешів: прибираємо, окрім самого префікса
        if path != '/' and path.endswith('/'):
            path = path[:-1]

        normalized = urlunparse((self.base_scheme, self.base_domain, path, '', '', ''))
        return normalized

    def is_valid_url(self, url):
        """Перевіряє, що URL внутрішній, у межах гілки contributing і не статичний."""
        if not url:
            return False
        parsed = urlparse(url)
        if parsed.scheme not in ('http', 'https'):
            return False
        if parsed.netloc != self.base_domain:
            return False
        if not parsed.path.startswith(self.base_path_prefix):
            return False
        if parsed.path.endswith(('.pdf', '.zip', '.tar', '.gz', '.md')):
            return False
        return True

    # =========================
    #  Витяг посилань і контенту
    # =========================
    def extract_links(self, soup, current_url):
        """Витягує внутрішні посилання у порядку появи (спершу навігація, далі контент)."""
        links = []
        seen = set()

        def add_link(href):
            full_url = urljoin(current_url, href)
            full_url = full_url.split('#')[0]
            if self.is_valid_url(full_url) and full_url not in seen:
                links.append(full_url)
                seen.add(full_url)

        # Посилання з навігації/сайдбару
        nav_menu = soup.find('nav') or soup.find('aside') or soup.find('div', class_=['sidebar', 'toctree'])
        if nav_menu:
            for a_tag in nav_menu.find_all('a', href=True):
                add_link(a_tag['href'])

        # Посилання з основного контенту
        main_content = (
            soup.find('main') or
            soup.find('article') or
            soup.find('div', class_='document') or
            soup.find('div', role='main') or
            soup.find('body')
        )
        if main_content:
            for a_tag in main_content.find_all('a', href=True):
                add_link(a_tag['href'])

        return links

    def extract_content(self, soup, url):
        """
        Витягує контент сторінки, очищаючи навігаційні та декоративні елементи.
        Повертає структуру: {'title': str, 'text': list[dict], 'url': str}
        """
        content_data = {'title': '', 'text': [], 'url': url}

        # Видаляємо небажані елементи
        for unwanted in soup.find_all(['nav', 'aside', 'header', 'footer', 'script', 'style']):
            unwanted.decompose()

        # Видаляємо за класами навігації/декору
        for unwanted_class in [
            'sidebar', 'navigation', 'breadcrumb', 'menu',
            'navbar', 'footer', 'header', 'btn', 'button', 'prev-next', 'toc', 'pager'
        ]:
            for element in soup.find_all(class_=lambda x: x and unwanted_class in str(x).lower()):
                element.decompose()

        main_content = (
            soup.find('main') or
            soup.find('article') or
            soup.find('div', class_='document') or
            soup.find('div', role='main') or
            soup.find('body')
        )

        if main_content:
            title_tag = main_content.find('h1')
            if title_tag:
                content_data['title'] = title_tag.get_text(strip=True)

            self._process_element(main_content, content_data['text'])

        return content_data

    def _process_element(self, element, text_list, parent_tag=None):
        """Рекурсивно витягує лише змістовний текст і структуру (включно з таблицями та кодом)."""
        if not hasattr(element, 'name') or element.name is None:
            return

        # Пропускаємо явні службові елементи
        if element.name in ['nav', 'aside', 'header', 'footer', 'button', 'script', 'style']:
            return

        # Пропускаємо елементи з навігаційними/декоративними класами
        element_class = element.get('class')
        if element_class:
            classes = ' '.join(element_class).lower()
            skip_classes = [
                'sidebar', 'breadcrumb', 'navigation', 'menu',
                'navbar', 'btn', 'button', 'prev-next', 'footer', 'header', 'toc', 'pager'
            ]
            if any(skip in classes for skip in skip_classes):
                return

        # Заголовки і параграфи
        if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'dt', 'dd']:
            # Уникаємо параграфів-збірок кнопок
            links = element.find_all('a')
            if links:
                all_buttons = all(
                    (link.get('class') and any(btn in ' '.join(link.get('class')).lower() for btn in ['btn', 'button']))
                    for link in links if link.get('class')
                )
                if all_buttons and len(links) > 0:
                    return

            text = element.get_text(" ", strip=True)
            if text and len(text) > 2:
                text_list.append({'type': element.name, 'text': text})
            # Не return — всередині заголовків інколи бувають інші елементи

        # Елементи списків
        if element.name == 'li':
            parent_ul = element.find_parent(['ul', 'ol'])
            if parent_ul:
                parent_class = parent_ul.get('class')
                if parent_class:
                    classes = ' '.join(parent_class).lower()
                    if any(skip in classes for skip in ['menu', 'nav', 'toc']):
                        return
            text = element.get_text(" ", strip=True)
            if text and len(text) > 2:
                text_list.append({'type': 'li', 'text': text})
            # Не виходимо рано — у li можуть бути вкладені <pre>/<code>/таблиці
            for child in element.children:
                if hasattr(child, 'name'):
                    self._process_element(child, text_list, element.name)
            return

        # Блоки коду: <pre> і/або <code>
        if element.name in ['pre', 'code']:
            # Якщо <code> всередині <pre> — обробляємо тільки <pre>
            if element.name == 'code' and element.find_parent('pre'):
                return

            # Визначення мови
            detected_lang = None
            if element.name == 'pre':
                dl = element.get('data-language')
                if dl:
                    detected_lang = str(dl).strip().lower()

            # language-*/lang-* на поточному елементі
            if not detected_lang:
                classes_here = element.get('class') or []
                if classes_here:
                    classes_joined = ' '.join(classes_here).lower()
                    for prefix in ('language-', 'lang-'):
                        if prefix in classes_joined:
                            for token in classes_joined.split():
                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]
                                    break
                            if detected_lang:
                                break

            # Якщо це <pre> і є <code class="language-..."> всередині
            if not detected_lang and element.name == 'pre':
                code_child = element.find('code')
                if code_child and code_child.get('class'):
                    cj = ' '.join(code_child.get('class')).lower()
                    for prefix in ('language-', 'lang-'):
                        if prefix in cj:
                            for token in cj.split():
                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]
                                    break
                            if detected_lang:
                                break

            # Випадок Sphinx: батьківський контейнер має клас виду "highlight-<lang>"
            if not detected_lang:
                parent_hl = element.find_parent('div', class_=lambda x: x and 'highlight-' in ' '.join(x).lower())
                if parent_hl and parent_hl.get('class'):
                    cj = ' '.join(parent_hl.get('class')).lower()
                    for token in cj.split():
                        if token.startswith('highlight-') and len(token) > len('highlight-'):
                            detected_lang = token[len('highlight-'):]
                            break

            # Витяг тексту коду зі збереженням рядків
            text = None
            if element.name == 'pre':
                # Формат Expressive Code: <div class="ec-line"><div class="code">...</div></div>
                ec_lines = element.select('div.ec-line div.code')
                if ec_lines:
                    lines = []
                    for code_div in ec_lines:
                        line_text = code_div.get_text('', strip=False)
                        lines.append(line_text.rstrip())
                    text = '\n'.join(lines).strip('\n')
                else:
                    # Якщо є внутрішній <code>
                    code_child = element.find('code')
                    if code_child:
                        # ВАЖЛИВО: не вставляємо штучні переносы між токенами
                        text = code_child.get_text('', strip=False).strip('\n')

            # Fallback: беремо текст елемента без агресивного strip і без \r
            if text is None:
                # ВАЖЛИВО: використати separator '' щоб не ламати відступи/пробіли між <span>
                text = element.get_text('', strip=False).replace('\r', '')

            # Нормалізація: прибираємо зайві порожні рядки з країв і пробіли у кінцях
            lines = [ln.rstrip() for ln in text.splitlines()]
            while lines and not lines[0].strip():
                lines.pop(0)
            while lines and not lines[-1].strip():
                lines.pop()
            text = '\n'.join(lines)

            # Опціонально: підпис до блоку коду (ім'я файлу) з Sphinx
            caption_text = None
            try:
                parent = element.parent
                if parent and getattr(parent, 'find', None):
                    cap_div = None
                    if 'literal-block-wrapper' in ' '.join((parent.get('class') or [])).lower():
                        cap_div = parent.find('div', class_=lambda x: x and 'code-block-caption' in ' '.join(x).lower())
                    if not cap_div:
                        cap_div = element.find_previous_sibling('div', class_=lambda x: x and 'code-block-caption' in ' '.join(x).lower())
                    if cap_div:
                        caption_text = cap_div.get_text(' ', strip=True)
                        if caption_text:
                            text_list.append({'type': 'code_caption', 'text': caption_text})
            except Exception:
                pass

            if text and len(text) > 2:
                text_list.append({'type': element.name, 'text': text, 'lang': detected_lang})
            return

        # Цитати
        if element.name == 'blockquote':
            text = element.get_text(" ", strip=True)
            if text and len(text) > 10:
                text_list.append({'type': 'blockquote', 'text': text})
            return

        # Таблиці
        if element.name == 'table':
            rows = []
            for tr in element.find_all('tr'):
                row = []
                for cell in tr.find_all(['th', 'td']):
                    row.append(cell.get_text(" ", strip=True))
                if any(row):
                    rows.append(row)
            if rows:
                text_list.append({'type': 'table', 'rows': rows})
            return

        # Рекурсія по вмістовних контейнерах
        if element.name in ['div', 'section', 'article', 'main', 'body', 'figure', 'details', 'ul', 'ol', 'dl']:
            for child in element.children:
                if hasattr(child, 'name'):
                    self._process_element(child, text_list, element.name)
            return

    # =========================
    #  Вивід у DOCX
    # =========================
    def add_to_document(self, content_data):
        """Записує сторінку у DOCX з форматуванням, уникненням дублювання h1 та підтримкою таблиць/коду."""
        # Навіть якщо немає текстових елементів, все одно додаємо URL і заголовок сторінки

        # Джерело
        p = self.doc.add_paragraph()
        run = p.add_run(f"Джерело: {content_data['url']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Заголовок сторінки (h1 додаємо один раз тут)
        if content_data['title']:
            heading = self.doc.add_heading(content_data['title'], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Контент
        for item in content_data['text']:
            t = item.get('type')
            # Уникаємо дублювання h1 (пропускаємо item з type == 'h1')
            if t in ['h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(t[1]) + 1
                self.doc.add_heading(item['text'], level=min(level, 9))

            elif t == 'code_caption':
                meta = self.doc.add_paragraph(item['text'])
                for run in meta.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(120, 120, 120)
                    run.font.italic = True

            elif t in ['pre', 'code']:
                lang = item.get('lang')
                if lang:
                    meta = self.doc.add_paragraph(f"Код ({lang})")
                    for run in meta.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(120, 120, 120)
                        run.font.italic = True

                p = self.doc.add_paragraph(item['text'])
                # Застосовуємо кастомний стиль коду, якщо доступний
                try:
                    p.style = self.doc.styles['Code Block']
                except KeyError:
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(50, 50, 50)

            elif t == 'blockquote':
                p = self.doc.add_paragraph(item['text'])
                try:
                    p.style = self.doc.styles['Quote Block']
                except KeyError:
                    p.paragraph_format.left_indent = Pt(36)
                    p.paragraph_format.right_indent = Pt(36)
                    for run in p.runs:
                        run.font.italic = True

            elif t == 'li':
                self.doc.add_paragraph(item['text'], style='List Bullet')

            elif t == 'table':
                rows = item.get('rows') or []
                if rows:
                    cols = max(len(r) for r in rows)
                    if cols > 0:
                        table = self.doc.add_table(rows=0, cols=cols)
                        # Встановлюємо стиль таблиці, якщо є
                        for style_name in ('Light List', 'Table Grid'):
                            try:
                                table.style = self.doc.styles[style_name]
                                break
                            except KeyError:
                                continue
                        for r in rows:
                            row_cells = table.add_row().cells
                            for i, cell_text in enumerate(r):
                                if i < len(row_cells):
                                    row_cells[i].text = cell_text

            else:
                txt = item.get('text', '')
                if txt and len(txt) > 5:
                    self.doc.add_paragraph(txt)

        # Роздільник між сторінками
        self.doc.add_paragraph('_' * 80)
        self.doc.add_paragraph()

    # =========================
    #  HTTP із ретраями
    # =========================
    def _get_soup(self, url):
        """Отримує BeautifulSoup з повторними спробами."""
        last_exc = None
        for attempt in range(self.retry + 1):
            try:
                resp = self.session.get(url, timeout=self.timeout)
                resp.raise_for_status()
                return BeautifulSoup(resp.content, 'html.parser')
            except Exception as e:
                last_exc = e
                # Якщо 404 і це посилання виду /contributing/<file>.html — пробуємо піддиректорії секцій
                try:
                    if hasattr(e, 'response') and e.response is not None and e.response.status_code == 404:
                        parsed = urlparse(url)
                        path = parsed.path
                        prefix = self.base_path_prefix + '/'
                        if path.startswith(prefix) and path.endswith('.html'):
                            tail = path[len(prefix):]  # '<file>.html' або з підшляхом
                            if tail and '/' not in tail:
                                file_name = tail.replace('.html', '')
                                for section in ('development', 'documentation'):
                                    alt_path = f"{self.base_path_prefix}/{section}/{file_name}.html"
                                    alt_url = urlunparse((parsed.scheme or self.base_scheme, parsed.netloc or self.base_domain, alt_path, '', '', ''))
                                    try:
                                        alt_resp = self.session.get(alt_url, timeout=self.timeout)
                                        alt_resp.raise_for_status()
                                        return BeautifulSoup(alt_resp.content, 'html.parser')
                                    except Exception:
                                        continue
                except Exception:
                    pass
                time.sleep(min(self.delay, 1.5))
        # Якщо всі спроби не вдалися, викидаємо останній виняток
        raise last_exc

    # =========================
    #  Збір структури URL
    # =========================
    def _collect_urls_recursive(self, url, depth=0):
        """Рекурсивно збирає URLs у порядку появи з уникненням дублікатів."""
        if url in self.visited:
            return

        self.visited.add(url)
        self.ordered_urls.append(url)

        if depth == 0:
            print(f"Сканування структури [{len(self.ordered_urls)}]: {url}")

        try:
            soup = self._get_soup(url)
            links = self.extract_links(soup, url)

            for link in links:
                if link not in self.visited:
                    if depth < 2:
                        print(f"Сканування структури [{len(self.ordered_urls)}]: {link}")
                    self._collect_urls_recursive(link, depth + 1)

            time.sleep(self.delay * 0.6)

        except requests.exceptions.HTTPError as e:
            print(f"Пропускаю {url}: {e}")
        except Exception as e:
            print(f"Помилка при скануванні {url}: {str(e)}")

    # =========================
    #  Оркестрація скрапінгу
    # =========================
    def scrape(self):
        """Основний процес: збір URL-ів та експорт контенту у встановленому порядку."""
        print(f"Початок скрапінгу: {self.start_url}")
        print("Спочатку збираємо всі URLs у правильному порядку...\n")

        self._collect_urls_recursive(self.start_url)

        total_pages = len(self.ordered_urls)
        print(f"\nЗнайдено {total_pages} сторінок")
        print(f"Починаємо завантаження контенту...\n")

        for idx, url in enumerate(self.ordered_urls, 1):
            try:
                print(f"Обробка [{idx}/{total_pages}]: {url}")
                soup = self._get_soup(url)
                content = self.extract_content(soup, url)
                self.add_to_document(content)
                time.sleep(self.delay)
            except requests.exceptions.HTTPError as e:
                print(f"Пропускаю {url}: {e}")
                continue
            except Exception as e:
                print(f"Помилка при обробці {url}: {str(e)}")
                continue

        print(f"\n✓ Завершено! Оброблено сторінок: {total_pages}")
        return total_pages

    def save(self, filename='odoo_18_contributing_docs.docx'):
        """Зберігає зібраний документ."""
        self.doc.save(filename)
        print(f"✓ Документ збережено: {filename}")

    # =========================
    #  Стилі DOCX
    # =========================
    def _ensure_styles(self):
        """
        Створює або налаштовує необхідні стилі DOCX:
        - Code Block: моноширинний стиль для коду
        - Quote Block: стиль для цитат
        """
        styles = self.doc.styles

        # Code Block
        try:
            code_style = styles['Code Block']
        except KeyError:
            try:
                code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                code_style = None
        if code_style is not None:
            try:
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(9)
                code_style.font.color.rgb = RGBColor(50, 50, 50)
            except Exception:
                pass

        # Quote Block
        try:
            quote_style = styles['Quote Block']
        except KeyError:
            try:
                quote_style = styles.add_style('Quote Block', WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                quote_style = None
        if quote_style is not None:
            try:
                quote_style.font.italic = True
                quote_style.font.color.rgb = RGBColor(80, 80, 80)
                pf = quote_style.paragraph_format
                pf.left_indent = Pt(36)
                pf.right_indent = Pt(36)
            except Exception:
                pass


def main():
    start_url = "https://www.odoo.com/documentation/18.0/contributing.html"
    scraper = Odoo18ContributingScraper(start_url, delay=0.8, retry=2, timeout=30)
    scraper.scrape()
    scraper.save('odoo_18_contributing_docs.docx')


if __name__ == "__main__":
    main()
