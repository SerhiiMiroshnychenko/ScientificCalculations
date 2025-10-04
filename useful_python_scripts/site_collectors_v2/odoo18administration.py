import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import time


class Odoo18AdministrationScraper:
    def __init__(self, start_url):
        # Базові параметри (логіка навігації 1:1 з оригіналом odoo18administration.py)
        self.start_url = start_url
        self.base_domain = urlparse(start_url).netloc
        self.base_path = '/'.join(urlparse(start_url).path.split('/')[:-1])
        self.visited = set()
        self.ordered_urls = []  # Порядок появи
        self.doc = Document()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        # Стилі для DOCX
        self._ensure_styles()

    # =========================
    #  Перевірка URL (як в оригіналі)
    # =========================
    def is_valid_url(self, url):
        """Перевірка чи URL належить до документації Odoo (administration)"""
        parsed = urlparse(url)
        # Перевірка домену
        if parsed.netloc != self.base_domain:
            return False
        # Перевірка що це сторінка потрібної гілки документації
        if not parsed.path.startswith('/documentation/18.0/administration'):
            return False
        # Виключаємо якорі та PDF файли
        if '#' in url or url.endswith('.pdf'):
            return False
        return True

    # =========================
    #  Витяг посилань (як в оригіналі)
    # =========================
    def extract_links(self, soup, current_url):
        """Витягує всі внутрішні посилання зі сторінки В ПОРЯДКУ появи в HTML"""
        links = []  # Використовуємо список замість set для збереження порядку
        seen = set()  # Для уникнення дублікатів

        # Спочатку шукаємо посилання в навігаційному меню (sidebar)
        nav_menu = soup.find('nav') or soup.find('aside') or soup.find('div', class_=['sidebar', 'toctree'])
        if nav_menu:
            # Витягуємо посилання з меню в порядку появи
            for a_tag in nav_menu.find_all('a', href=True):
                href = a_tag['href']
                full_url = urljoin(current_url, href)
                full_url = full_url.split('#')[0]
                if self.is_valid_url(full_url) and full_url not in seen:
                    links.append(full_url)
                    seen.add(full_url)

        # Потім додаємо посилання з основного контенту
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='document')
        if main_content:
            for a_tag in main_content.find_all('a', href=True):
                href = a_tag['href']
                full_url = urljoin(current_url, href)
                full_url = full_url.split('#')[0]
                if self.is_valid_url(full_url) and full_url not in seen:
                    links.append(full_url)
                    seen.add(full_url)

        return links

    # =========================
    #  Витяг контенту (оригінальна база + покращення)
    # =========================
    def extract_content(self, soup, url):
        """Витягує основний контент зі сторінки та структурує його"""
        content_data = {'title': '', 'text': [], 'url': url}

        # Видаляємо небажані елементи перед обробкою
        for unwanted in soup.find_all(['nav', 'aside', 'header', 'footer', 'script', 'style']):
            unwanted.decompose()
        # Видаляємо елементи навігації за класами (як в оригіналі)
        for unwanted_class in ['sidebar', 'navigation', 'toctree', 'breadcrumb', 'menu',
                               'navbar', 'footer', 'header', 'btn', 'button']:
            for element in soup.find_all(class_=lambda x: x and unwanted_class in str(x).lower()):
                element.decompose()

        # Знаходимо основний контент
        main_content = (soup.find('main') or
                        soup.find('article') or
                        soup.find('div', class_='document') or
                        soup.find('div', role='main'))
        if not main_content:
            main_content = soup.find('body')

        if main_content:
            # Знаходимо заголовок статті
            title_tag = main_content.find('h1')
            if title_tag:
                content_data['title'] = title_tag.get_text(strip=True)
            # Обробляємо контент рекурсивно, зберігаючи структуру
            self._process_element(main_content, content_data['text'])

        return content_data

    def _process_element(self, element, text_list, parent_tag=None):
        """Рекурсивно обробляє елементи, витягуючи тільки змістовний контент (з таблицями та кодом)"""
        # Перевіряємо чи це взагалі HTML елемент
        if not hasattr(element, 'name') or element.name is None:
            return

        # Пропускаємо елементи навігації та кнопки
        if element.name in ['nav', 'aside', 'header', 'footer', 'button', 'script', 'style']:
            return

        # Пропускаємо елементи з класами навігації
        element_class = element.get('class')
        if element_class:
            classes = ' '.join(element_class).lower()
            skip_classes = ['sidebar', 'toctree', 'breadcrumb', 'navigation', 'menu',
                            'navbar', 'btn', 'button', 'prev-next', 'footer', 'header']
            if any(skip in classes for skip in skip_classes):
                return

        # Обробляємо текстові блоки (додано dt/dd)
        if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'dt', 'dd']:
            # Перевіряємо чи є в параграфі посилання-кнопки
            links = element.find_all('a')
            if links:
                # Якщо всі посилання мають клас кнопки - пропускаємо
                all_buttons = all(
                    link.get('class') and any(btn in ' '.join(link.get('class')).lower() for btn in ['btn', 'button'])
                    for link in links if link.get('class')
                )
                if all_buttons and len(links) > 0:
                    return
            text = element.get_text(' ', strip=True)
            # Фільтруємо короткі тексти які можуть бути частиною навігації
            if text and len(text) > 2:
                text_list.append({'type': element.name, 'text': text})

        # Обробляємо списки (не втрачаємо вкладені pre/code/table)
        elif element.name == 'li':
            # Пропускаємо елементи списку в навігації
            parent_ul = element.find_parent(['ul', 'ol'])
            if parent_ul:
                parent_class = parent_ul.get('class')
                if parent_class:
                    classes = ' '.join(parent_class).lower()
                    if any(skip in classes for skip in ['toctree', 'menu', 'nav']):
                        return
            text = element.get_text(' ', strip=True)
            if text and len(text) > 2:
                text_list.append({'type': 'li', 'text': text})
            # Рекурсія по дітях
            for child in element.children:
                if hasattr(child, 'name'):
                    self._process_element(child, text_list, element.name)

        # Обробляємо блоки коду — ПОКРАЩЕННЯ
        elif element.name in ['pre', 'code']:
            # Для code всередині pre - пропускаємо (обробимо pre)
            if element.name == 'code' and element.find_parent('pre'):
                return

            # Визначення мови
            detected_lang = None
            if element.name == 'pre':
                dl = element.get('data-language')
                if dl:
                    detected_lang = str(dl).strip().lower()
            if not detected_lang:
                classes_here = element.get('class') or []
                if classes_here:
                    cj = ' '.join(classes_here).lower()
                    for prefix in ('language-', 'lang-'):
                        if prefix in cj:
                            for token in cj.split():
                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]
                                    break
                            if detected_lang:
                                break
            if not detected_lang and element.name == 'pre':
                code_child = element.find('code')
                if code_child and code_child.get('class'):
                    cj = ' '.join(code_child.get('class')).lower()
                    for prefix in ('language-', 'lang-'):
                        if prefix in cj:
                            for token in cj.split():
                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]
                                    break
                            if detected_lang:
                                break
            if not detected_lang:
                parent_hl = element.find_parent('div', class_=lambda x: x and 'highlight-' in ' '.join(x).lower())
                if parent_hl and parent_hl.get('class'):
                    cj = ' '.join(parent_hl.get('class')).lower()
                    for token in cj.split():
                        if token.startswith('highlight-') and len(token) > len('highlight-'):
                            detected_lang = token[len('highlight-'):]
                            break

            # Текст коду зі збереженням рядків та відступів
            text = None
            if element.name == 'pre':
                # Формат Expressive Code: <div class="ec-line"><div class="code">...</div></div>
                ec_lines = element.select('div.ec-line div.code')
                if ec_lines:
                    lines = []
                    for code_div in ec_lines:
                        line_text = code_div.get_text('', strip=False)
                        lines.append(line_text.rstrip())
                    text = '\n'.join(lines).strip('\n')
                else:
                    # Якщо є внутрішній <code>
                    code_child = element.find('code')
                    if code_child:
                        text = code_child.get_text('', strip=False).strip('\n')
            # Fallback: беремо текст елемента без агресивного strip і без \r
            if text is None:
                text = element.get_text('', strip=False).replace('\r', '')

            # Нормалізація: прибираємо зайві порожні рядки з країв і пробіли у кінцях
            lines = [ln.rstrip() for ln in text.splitlines()]
            while lines and not lines[0].strip():
                lines.pop(0)
            while lines and not lines[-1].strip():
                lines.pop()
            text = '\n'.join(lines)

            # Підпис до коду (Sphinx): code-block-caption
            try:
                parent = element.parent
                if parent and getattr(parent, 'find', None):
                    cap_div = None
                    if 'literal-block-wrapper' in ' '.join((parent.get('class') or [])).lower():
                        cap_div = parent.find('div', class_=lambda x: x and 'code-block-caption' in ' '.join(x).lower())
                    if not cap_div:
                        cap_div = element.find_previous_sibling('div', class_=lambda x: x and 'code-block-caption' in ' '.join(x).lower())
                    if cap_div:
                        caption_text = cap_div.get_text(' ', strip=True)
                        if caption_text:
                            text_list.append({'type': 'code_caption', 'text': caption_text})
            except Exception:
                pass

            if text and len(text) > 2:
                text_list.append({'type': element.name, 'text': text, 'lang': detected_lang})

        # Обробляємо blockquote
        elif element.name == 'blockquote':
            text = element.get_text(' ', strip=True)
            if text and len(text) > 10:
                text_list.append({'type': 'blockquote', 'text': text})

        # Таблиці — ПОКРАЩЕННЯ
        elif element.name == 'table':
            rows = []
            for tr in element.find_all('tr'):
                row = []
                for cell in tr.find_all(['th', 'td']):
                    row.append(cell.get_text(' ', strip=True))
                if any(row):
                    rows.append(row)
            if rows:
                text_list.append({'type': 'table', 'rows': rows})

        # Рекурсивно обробляємо дочірні елементи для div, section тощо (розширено)
        elif element.name in ['div', 'section', 'article', 'main', 'body', 'figure', 'details', 'ul', 'ol', 'dl']:
            for child in element.children:
                if hasattr(child, 'name'):
                    self._process_element(child, text_list, element.name)

    # =========================
    #  Додає контент у DOCX (покращено стилі та недопущення дублювання h1)
    # =========================
    def add_to_document(self, content_data):
        """Додає контент до DOCX документу"""
        if not content_data['text']:
            # Все одно додамо джерело та заголовок, аби верхньорівневі індексні сторінки були присутні
            p = self.doc.add_paragraph()
            run = p.add_run(f"Джерело: {content_data['url']}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
            if content_data.get('title'):
                heading = self.doc.add_heading(content_data['title'], level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.doc.add_paragraph('_' * 80)
            self.doc.add_paragraph()
            return

        # Додаємо URL як коментар
        p = self.doc.add_paragraph()
        run = p.add_run(f"Джерело: {content_data['url']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Додаємо заголовок сторінки (h1 додається лише тут)
        if content_data['title']:
            heading = self.doc.add_heading(content_data['title'], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Додаємо текстовий контент
        for item in content_data['text']:
            t = item.get('type')
            if t in ['h2', 'h3', 'h4']:
                level = int(t[1]) + 1  # h2 -> level 3, h3 -> level 4, etc.
                self.doc.add_heading(item['text'], level=min(level, 9))
            elif t == 'code_caption':
                meta = self.doc.add_paragraph(item['text'])
                for r in meta.runs:
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(120, 120, 120)
                    r.font.italic = True
            elif t in ['pre', 'code']:
                lang = item.get('lang')
                if lang:
                    meta = self.doc.add_paragraph(f"Код ({lang})")
                    for r in meta.runs:
                        r.font.size = Pt(8)
                        r.font.color.rgb = RGBColor(120, 120, 120)
                        r.font.italic = True
                p = self.doc.add_paragraph(item['text'])
                try:
                    p.style = self.doc.styles['Code Block']
                except KeyError:
                    p.style = 'No Spacing'
                    for r in p.runs:
                        r.font.name = 'Courier New'
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(50, 50, 50)
            elif t == 'blockquote':
                p = self.doc.add_paragraph(item['text'])
                try:
                    p.style = self.doc.styles['Quote Block']
                except KeyError:
                    p.paragraph_format.left_indent = Pt(36)
                    p.paragraph_format.right_indent = Pt(36)
                    for r in p.runs:
                        r.font.italic = True
            elif t == 'li':
                self.doc.add_paragraph(item['text'], style='List Bullet')
            elif t == 'table':
                rows = item.get('rows') or []
                if rows:
                    cols = max(len(r) for r in rows)
                    if cols > 0:
                        table = self.doc.add_table(rows=0, cols=cols)
                        for style_name in ('Light List', 'Table Grid'):
                            try:
                                table.style = self.doc.styles[style_name]
                                break
                            except KeyError:
                                continue
                        for r in rows:
                            row_cells = table.add_row().cells
                            for i, cell_text in enumerate(r):
                                if i < len(row_cells):
                                    row_cells[i].text = cell_text
            else:
                txt = item.get('text', '')
                if txt and len(txt) > 5:  # Фільтруємо дуже короткі тексти
                    self.doc.add_paragraph(txt)

        # Додаємо роздільник
        self.doc.add_paragraph('_' * 80)
        self.doc.add_paragraph()

    # =========================
    #  Основний метод скрапінгу (як в оригіналі)
    # =========================
    def scrape(self):
        print(f"Початок скрапінгу: {self.start_url}")
        print(f"Спочатку збираємо всі URLs в правильному порядку...\n")
        # Фаза 1: Збираємо всі URLs в правильному порядку
        self._collect_urls_recursive(self.start_url)

        total_pages = len(self.ordered_urls)
        print(f"\nЗнайдено {total_pages} сторінок")
        print(f"Починаємо завантаження контенту...\n")

        # Фаза 2: Завантажуємо контент в зібраному порядку
        for idx, url in enumerate(self.ordered_urls, 1):
            try:
                print(f"Обробка [{idx}/{total_pages}]: {url}")
                response = self.session.get(url, timeout=30)
                response.raise_for_status()
                soup = BeautifulSoup(response.content, 'html.parser')
                content = self.extract_content(soup, url)
                self.add_to_document(content)
                # Затримка щоб не перевантажувати сервер
                time.sleep(1)
            except Exception as e:
                print(f"Помилка при обробці {url}: {str(e)}")
                continue

        print(f"\n✓ Завершено! Оброблено сторінок: {total_pages}")
        return total_pages

    def _collect_urls_recursive(self, url, depth=0):
        """Рекурсивно збирає URLs в порядку появи в навігації"""
        if url in self.visited or depth > 50:  # Обмеження глибини
            return
        self.visited.add(url)
        self.ordered_urls.append(url)
        if depth == 0:
            print(f"Сканування структури [{len(self.ordered_urls)}]: {url}")
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            # Витягуємо посилання в порядку появи
            links = self.extract_links(soup, url)
            # Рекурсивно обходимо кожне посилання
            for link in links:
                if link not in self.visited:
                    if depth < 2:  # Показуємо прогрес тільки для перших рівнів
                        print(f"Сканування структури [{len(self.ordered_urls)}]: {link}")
                    self._collect_urls_recursive(link, depth + 1)
            time.sleep(0.5)  # Коротша затримка при скануванні структури
        except Exception as e:
            print(f"Помилка при скануванні {url}: {str(e)}")

    def save(self, filename='odoo_18_administration_docs.docx'):
        """Зберігає документ"""
        self.doc.save(filename)
        print(f"✓ Документ збережено: {filename}")

    # =========================
    #  Стилі DOCX (покращення)
    # =========================
    def _ensure_styles(self):
        styles = self.doc.styles
        # Code Block
        try:
            code_style = styles['Code Block']
        except KeyError:
            try:
                code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                code_style = None
        if code_style is not None:
            try:
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(9)
                code_style.font.color.rgb = RGBColor(50, 50, 50)
            except Exception:
                pass
        # Quote Block
        try:
            quote_style = styles['Quote Block']
        except KeyError:
            try:
                quote_style = styles.add_style('Quote Block', WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                quote_style = None
        if quote_style is not None:
            try:
                quote_style.font.italic = True
                quote_style.font.color.rgb = RGBColor(80, 80, 80)
                pf = quote_style.paragraph_format
                pf.left_indent = Pt(36)
                pf.right_indent = Pt(36)
            except Exception:
                pass


def main():
    start_url = "https://www.odoo.com/documentation/18.0/administration.html"
    scraper = Odoo18AdministrationScraper(start_url)
    scraper.scrape()
    scraper.save('odoo_18_administration_docs.docx')


if __name__ == "__main__":
    main()
