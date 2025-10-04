import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, urlunparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import time


class OwlDocsScraper:
    def __init__(self, start_url, delay=0.8, retry=2, timeout=30):
        self.start_url = start_url.rstrip('/') + '/'
        parsed = urlparse(self.start_url)
        self.base_scheme = parsed.scheme or 'https'
        self.base_domain = parsed.netloc
        self.visited = set()
        self.ordered_urls = []
        self.doc = Document()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        self.delay = delay
        self.retry = retry
        self.timeout = timeout
        # Перевіряємо/створюємо необхідні стилі для DOCX
        self._ensure_styles()

    def normalize_url(self, url, current_url=None):
        """Нормалізує URL: абсолютний шлях у межах домену, без якорів та параметрів."""
        if not url:
            return None

        low = url.strip().lower()
        if low.startswith(('mailto:', 'javascript:', 'tel:')):
            return None

        base = current_url or self.start_url
        absolute = urljoin(base, url)

        parsed = urlparse(absolute)
        cleaned = parsed._replace(fragment='', query='')

        if cleaned.netloc != self.base_domain:
            return None

        path = cleaned.path or '/'
        if not path.startswith('/'):
            path = '/' + path
        # Знімаємо трейлінг-слеш для узгодженості (крім кореня)
        if path != '/' and path.endswith('/'):
            path = path[:-1]

        # Ігноруємо статичні файли
        if path.endswith(('.pdf', '.zip', '.tar', '.gz', '.md')):
            return None

        normalized = urlunparse((self.base_scheme, self.base_domain, path, '', '', ''))
        return normalized

    def is_valid_url(self, url):
        """Перевіряє, що URL внутрішній та не веде на статичні файли."""
        if not url:
            return False
        parsed = urlparse(url)
        if parsed.scheme not in ('http', 'https'):
            return False
        if parsed.netloc != self.base_domain:
            return False
        if parsed.path.endswith(('.pdf', '.zip', '.tar', '.gz', '.md')):
            return False
        return True

    def extract_links(self, soup, current_url):
        """Витягує всі внутрішні посилання зі сторінки в порядку появи."""
        links = []
        seen = set()

        def add_link(href):
            norm = self.normalize_url(href, current_url=current_url)
            if not norm or not self.is_valid_url(norm):
                return
            if norm in seen:
                return
            seen.add(norm)
            links.append(norm)

        # Посилання з навігації/сайдбару
        nav_menu = soup.find('nav') or soup.find('aside') or soup.find('div', class_=['sidebar', 'toctree'])
        if nav_menu:
            for a_tag in nav_menu.find_all('a', href=True):
                add_link(a_tag['href'])

        # Посилання з основного контенту
        main_content = (
            soup.find('main') or
            soup.find('article') or
            soup.find('div', class_='document') or
            soup.find('div', role='main') or
            soup.find('body')
        )
        if main_content:
            for a_tag in main_content.find_all('a', href=True):
                add_link(a_tag['href'])

        return links

    def extract_content(self, soup, url):
        """Витягує контент сторінки, очищаючи навігацію та декоративні елементи."""
        content_data = {'title': '', 'text': [], 'url': url}

        # Видаляємо небажані елементи
        for unwanted in soup.find_all(['nav', 'aside', 'header', 'footer', 'script', 'style']):
            unwanted.decompose()

        # Видаляємо за класами навігації/декору
        for unwanted_class in [
            'sidebar', 'navigation', 'toctree', 'breadcrumb', 'menu',
            'navbar', 'footer', 'header', 'btn', 'button', 'prev-next', 'toc', 'pager'
        ]:
            for element in soup.find_all(class_=lambda x: x and unwanted_class in str(x).lower()):
                element.decompose()

        main_content = (
            soup.find('main') or
            soup.find('article') or
            soup.find('div', class_='document') or
            soup.find('div', role='main') or
            soup.find('body')
        )

        if main_content:
            title_tag = main_content.find('h1')
            if title_tag:
                content_data['title'] = title_tag.get_text(strip=True)

            self._process_element(main_content, content_data['text'])

        return content_data

    def _process_element(self, element, text_list, parent_tag=None):
        """Рекурсивно витягує лише змістовний текст."""
        if not hasattr(element, 'name') or element.name is None:
            return

        if element.name in ['nav', 'aside', 'header', 'footer', 'button', 'script', 'style']:
            return

        element_class = element.get('class')
        if element_class:
            classes = ' '.join(element_class).lower()
            skip_classes = [
                'sidebar', 'toctree', 'breadcrumb', 'navigation', 'menu',
                'navbar', 'btn', 'button', 'prev-next', 'footer', 'header', 'toc', 'pager'
            ]
            if any(skip in classes for skip in skip_classes):
                return

        if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # уникаємо "кнопкових" посилань
            links = element.find_all('a')
            if links:
                all_buttons = all(
                    (link.get('class') and any(btn in ' '.join(link.get('class')).lower() for btn in ['btn', 'button']))
                    for link in links if link.get('class')
                )
                if all_buttons and len(links) > 0:
                    return

            text = element.get_text(" ", strip=True)
            if text and len(text) > 2:
                text_list.append({'type': element.name, 'text': text})

        elif element.name == 'li':
            parent_ul = element.find_parent(['ul', 'ol'])
            if parent_ul:
                parent_class = parent_ul.get('class')
                if parent_class:
                    classes = ' '.join(parent_class).lower()
                    if any(skip in classes for skip in ['toctree', 'menu', 'nav', 'toc']):
                        return
            text = element.get_text(" ", strip=True)
            if text and len(text) > 2:
                text_list.append({'type': 'li', 'text': text})




        elif element.name in ['pre', 'code']:

            # якщо <code> всередині <pre>, обробляємо тільки <pre>

            if element.name == 'code' and element.find_parent('pre'):
                return

            # 1) Визначення мови

            detected_lang = None

            # 1.1) <pre data-language="...">

            if element.name == 'pre':

                dl = element.get('data-language')

                if dl:
                    detected_lang = str(dl).strip().lower()

            # 1.2) Класи language-*/lang-* на поточному елементі

            if not detected_lang:

                classes_here = element.get('class') or []

                if classes_here:

                    classes_joined = ' '.join(classes_here).lower()

                    for prefix in ('language-', 'lang-'):

                        if prefix in classes_joined:

                            for token in classes_joined.split():

                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]

                                    break

                            if detected_lang:
                                break

            # 1.3) Якщо це <pre> і є <code class="language-..."> всередині

            if not detected_lang and element.name == 'pre':

                code_child = element.find('code')

                if code_child and code_child.get('class'):

                    cj = ' '.join(code_child.get('class')).lower()

                    for prefix in ('language-', 'lang-'):

                        if prefix in cj:

                            for token in cj.split():

                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]

                                    break

                            if detected_lang:
                                break

            # 2) Акуратний витяг тексту коду зі збереженням рядків (виконуємо ЗАВЖДИ)

            text = None

            if element.name == 'pre':

                # 2.1) Формат Expressive Code: <div class="ec-line"><div class="code">...</div></div>

                ec_lines = element.select('div.ec-line div.code')

                if ec_lines:

                    lines = []

                    for code_div in ec_lines:
                        # Беремо текст рядка без вставлення зайвих пробілів/переносів

                        line_text = code_div.get_text('', strip=False)

                        lines.append(line_text.rstrip())

                    text = '\n'.join(lines).strip('\n')

                else:

                    # 2.2) Якщо є внутрішній <code> — беремо його як суцільний текст

                    code_child = element.find('code')

                    if code_child:
                        text = code_child.get_text('', strip=False).strip('\n')

            # 2.3) Загальний fallback: беремо текст елемента без агресивного strip і без \r

            if text is None:
                text = element.get_text('\n', strip=False).replace('\r', '')

            # 2.4) Нормалізація: прибираємо зайві пробіли в кінці рядків і порожні крайні рядки

            lines = [ln.rstrip() for ln in text.splitlines()]

            while lines and not lines[0].strip():
                lines.pop(0)

            while lines and not lines[-1].strip():
                lines.pop()

            text = '\n'.join(lines)

            if text and len(text) > 2:
                text_list.append({'type': element.name, 'text': text, 'lang': detected_lang})


            if not detected_lang and element.name == 'pre':

                code_child = element.find('code')

                if code_child and code_child.get('class'):

                    cj = ' '.join(code_child.get('class')).lower()

                    for prefix in ('language-', 'lang-'):

                        if prefix in cj:

                            for token in cj.split():

                                if token.startswith(prefix) and len(token) > len(prefix):
                                    detected_lang = token[len(prefix):]

                                    break

                            if detected_lang:
                                break

                # Акуратний витяг тексту коду зі збереженням рядків
                text = None

                if element.name == 'pre':
                    # 1) Пробуємо формат Expressive Code: <div class="ec-line"><div class="code">...</div></div>
                    ec_lines = element.select('div.ec-line div.code')
                    if ec_lines:
                        lines = []
                        for code_div in ec_lines:
                            # Беремо текст рядка без додаткових розділювачів між <span>, щоб не ламати пробіли
                            line_text = code_div.get_text('', strip=False)
                            lines.append(line_text.rstrip())
                        text = '\n'.join(lines).strip('\n')
                    else:
                        # 2) Якщо є внутрішній <code> — беремо його як суцільний текст
                        code_child = element.find('code')
                        if code_child:
                            text = code_child.get_text('', strip=False).strip('\n')

                # 3) Загальний fallback: беремо текст елемента без агресивного strip і без \r
                if text is None:
                    text = element.get_text('\n', strip=False).replace('\r', '')

                # 4) Нормалізація: прибираємо зайві пробіли в кінці рядків і порожні “шумові” рядки
                lines = [ln.rstrip() for ln in text.splitlines()]
                # Прибрати провідні/завершальні порожні лінії
                while lines and not lines[0].strip():
                    lines.pop(0)
                while lines and not lines[-1].strip():
                    lines.pop()

                text = '\n'.join(lines)

                if text and len(text) > 2:
                    text_list.append({'type': element.name, 'text': text, 'lang': detected_lang})

        elif element.name == 'blockquote':
            text = element.get_text(" ", strip=True)
            if text and len(text) > 10:
                text_list.append({'type': 'blockquote', 'text': text})


        elif element.name == 'table':

            # Перенесення таблиці як рядків і комірок

            rows = []

            for tr in element.find_all('tr'):

                row = []

                for cell in tr.find_all(['th', 'td']):
                    row.append(cell.get_text(" ", strip=True))

                if any(row):
                    rows.append(row)

            if rows:
                text_list.append({'type': 'table', 'rows': rows})


        elif element.name in ['div', 'section', 'article', 'main', 'body', 'figure', 'details']:

            for child in element.children:

                if hasattr(child, 'name'):
                    self._process_element(child, text_list, element.name)

    def add_to_document(self, content_data):
        """Записує сторінку у DOCX з легким форматуванням."""
        if not content_data['text']:
            return

        # Джерело
        p = self.doc.add_paragraph()
        run = p.add_run(f"Джерело: {content_data['url']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Заголовок сторінки
        if content_data['title']:
            heading = self.doc.add_heading(content_data['title'], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Контент
        for item in content_data['text']:
            # Уникаємо дублювання h1 (заголовок сторінки вже додано вище)
            if item['type'] in ['h2', 'h3', 'h4']:
                level = int(item['type'][1]) + 1
                self.doc.add_heading(item['text'], level=min(level, 9))

            elif item['type'] in ['pre', 'code']:
                # Невеликий підпис з мовою, якщо визначено
                lang = item.get('lang')
                if lang:
                    meta = self.doc.add_paragraph(f"Код ({lang})")
                    for run in meta.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(120, 120, 120)
                        run.font.italic = True

                p = self.doc.add_paragraph(item['text'])
                # Спроба застосувати кастомний стиль для коду
                try:
                    p.style = self.doc.styles['Code Block']
                except KeyError:
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(50, 50, 50)

            elif item['type'] == 'blockquote':
                p = self.doc.add_paragraph(item['text'])
                # Спроба застосувати кастомний стиль для цитат
                try:
                    p.style = self.doc.styles['Quote Block']
                except KeyError:
                    p.paragraph_format.left_indent = Pt(36)
                    p.paragraph_format.right_indent = Pt(36)
                    for run in p.runs:
                        run.font.italic = True

            elif item['type'] == 'li':
                self.doc.add_paragraph(item['text'], style='List Bullet')

            elif item.get('type') == 'table':
                rows = item.get('rows') or []
                if rows:
                    cols = max(len(r) for r in rows)
                    if cols > 0:
                        table = self.doc.add_table(rows=0, cols=cols)
                        # Спробуємо встановити відомий стиль таблиці
                        for style_name in ('Light List', 'Table Grid'):
                            try:
                                table.style = self.doc.styles[style_name]
                                break
                            except KeyError:
                                continue
                        for r in rows:
                            row_cells = table.add_row().cells
                            for i, cell_text in enumerate(r):
                                if i < len(row_cells):
                                    row_cells[i].text = cell_text

            else:
                if len(item['text']) > 5:
                    self.doc.add_paragraph(item['text'])

        # Роздільник між сторінками
        self.doc.add_paragraph('_' * 80)
        self.doc.add_paragraph()

    def _get_soup(self, url):
        """Отримує BeautifulSoup з повторними спробами."""
        last_exc = None
        for attempt in range(self.retry + 1):
            try:
                resp = self.session.get(url, timeout=self.timeout)
                resp.raise_for_status()
                return BeautifulSoup(resp.content, 'html.parser')
            except Exception as e:
                last_exc = e
                time.sleep(min(self.delay, 1.5))
        # Якщо всі спроби не вдалися, викидаємо останній виняток
        raise last_exc

    def _collect_urls_recursive(self, url, depth=0):
        """Рекурсивно збирає URLs у порядку появи з уникненням дублікатів."""
        if url in self.visited:
            return

        self.visited.add(url)
        self.ordered_urls.append(url)

        if depth == 0:
            print(f"Сканування структури [{len(self.ordered_urls)}]: {url}")

        try:
            soup = self._get_soup(url)
            links = self.extract_links(soup, url)

            for link in links:
                if link not in self.visited:
                    if depth < 2:
                        print(f"Сканування структури [{len(self.ordered_urls)}]: {link}")
                    self._collect_urls_recursive(link, depth + 1)

            time.sleep(self.delay * 0.6)

        except requests.exceptions.HTTPError as e:
            print(f"Пропускаю {url}: {e}")
        except Exception as e:
            print(f"Помилка при скануванні {url}: {str(e)}")

    def scrape(self):
        """Основний процес: збір URL-ів та експорт контенту у встановленому порядку."""
        print(f"Початок скрапінгу: {self.start_url}")
        print("Спочатку збираємо всі URLs у правильному порядку...\n")

        self._collect_urls_recursive(self.start_url)

        total_pages = len(self.ordered_urls)
        print(f"\nЗнайдено {total_pages} сторінок")
        print(f"Починаємо завантаження контенту...\n")

        for idx, url in enumerate(self.ordered_urls, 1):
            try:
                print(f"Обробка [{idx}/{total_pages}]: {url}")
                soup = self._get_soup(url)
                content = self.extract_content(soup, url)
                self.add_to_document(content)
                time.sleep(self.delay)
            except requests.exceptions.HTTPError as e:
                print(f"Пропускаю {url}: {e}")
                continue
            except Exception as e:
                print(f"Помилка при обробці {url}: {str(e)}")
                continue

        print(f"\n✓ Завершено! Оброблено сторінок: {total_pages}")
        return total_pages

    def save(self, filename='owl-docs.docx'):
        """Зберігає зібраний документ."""
        self.doc.save(filename)
        print(f"✓ Документ збережено: {filename}")

    def _ensure_styles(self):
        """Створює або налаштовує необхідні стилі DOCX.
        - Code Block: моноширинний стиль для коду
        - Quote Block: стиль для цитат
        """
        styles = self.doc.styles

        # Code Block
        try:
            code_style = styles['Code Block']
        except KeyError:
            try:
                code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                code_style = None
        if code_style is not None:
            try:
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(9)
                code_style.font.color.rgb = RGBColor(50, 50, 50)
            except Exception:
                pass

        # Quote Block
        try:
            quote_style = styles['Quote Block']
        except KeyError:
            try:
                quote_style = styles.add_style('Quote Block', WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                quote_style = None
        if quote_style is not None:
            try:
                quote_style.font.italic = True
                quote_style.font.color.rgb = RGBColor(80, 80, 80)
                pf = quote_style.paragraph_format
                pf.left_indent = Pt(36)
                pf.right_indent = Pt(36)
            except Exception:
                pass


def main():
    start_url = "https://owl-docs.vercel.app/"
    scraper = OwlDocsScraper(start_url, delay=0.8, retry=2, timeout=30)
    scraper.scrape()
    scraper.save('owl-docs.docx')


if __name__ == "__main__":
    main()