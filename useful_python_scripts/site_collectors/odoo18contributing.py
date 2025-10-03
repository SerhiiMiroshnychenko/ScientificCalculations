import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
from collections import deque


class OdooDocsScraper:
    def __init__(self, start_url):
        self.start_url = start_url
        self.base_domain = urlparse(start_url).netloc
        self.base_path = '/'.join(urlparse(start_url).path.split('/')[:-1])
        self.visited = set()
        self.ordered_urls = []  # Зберігаємо URLs в порядку знаходження
        self.doc = Document()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })

    def is_valid_url(self, url):
        """Перевірка чи URL належить до документації Odoo"""
        parsed = urlparse(url)

        # Перевірка домену
        if parsed.netloc != self.base_domain:
            return False

        # Перевірка що це сторінка документації
        if not parsed.path.startswith('/documentation/18.0/contributing'):
            return False

        # Виключаємо якорі та PDF файли
        if '#' in url or url.endswith('.pdf'):
            return False

        return True

    def extract_links(self, soup, current_url):
        """Витягує всі внутрішні посилання зі сторінки В ПОРЯДКУ появи в HTML"""
        links = []  # Використовуємо список замість set для збереження порядку
        seen = set()  # Для уникнення дублікатів

        # Спочатку шукаємо посилання в навігаційному меню (sidebar)
        nav_menu = soup.find('nav') or soup.find('aside') or soup.find('div', class_=['sidebar', 'toctree'])

        if nav_menu:
            # Витягуємо посилання з меню в порядку появи
            for a_tag in nav_menu.find_all('a', href=True):
                href = a_tag['href']
                full_url = urljoin(current_url, href)
                full_url = full_url.split('#')[0]

                if self.is_valid_url(full_url) and full_url not in seen:
                    links.append(full_url)
                    seen.add(full_url)

        # Потім додаємо посилання з основного контенту
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='document')

        if main_content:
            for a_tag in main_content.find_all('a', href=True):
                href = a_tag['href']
                full_url = urljoin(current_url, href)
                full_url = full_url.split('#')[0]

                if self.is_valid_url(full_url) and full_url not in seen:
                    links.append(full_url)
                    seen.add(full_url)

        return links

    def extract_content(self, soup, url):
        """Витягує основний контент зі сторінки"""
        content_data = {
            'title': '',
            'text': [],
            'url': url
        }

        # Видаляємо небажані елементи перед обробкою
        for unwanted in soup.find_all(['nav', 'aside', 'header', 'footer', 'script', 'style']):
            unwanted.decompose()

        # Видаляємо елементи навігації за класами
        for unwanted_class in ['sidebar', 'navigation', 'toctree', 'breadcrumb', 'menu',
                               'navbar', 'footer', 'header', 'btn', 'button']:
            for element in soup.find_all(class_=lambda x: x and unwanted_class in str(x).lower()):
                element.decompose()

        # Знаходимо основний контент
        main_content = (soup.find('main') or
                        soup.find('article') or
                        soup.find('div', class_='document') or
                        soup.find('div', role='main'))

        if not main_content:
            main_content = soup.find('body')

        if main_content:
            # Знаходимо заголовок статті
            title_tag = main_content.find('h1')
            if title_tag:
                content_data['title'] = title_tag.get_text(strip=True)

            # Обробляємо контент рекурсивно, зберігаючи структуру
            self._process_element(main_content, content_data['text'])

        return content_data

    def _process_element(self, element, text_list, parent_tag=None):
        """Рекурсивно обробляє елементи, витягуючи тільки текстовий контент"""

        # Перевіряємо чи це взагалі HTML елемент
        if not hasattr(element, 'name') or element.name is None:
            return

        # Пропускаємо елементи навігації та кнопки
        if element.name in ['nav', 'aside', 'header', 'footer', 'button', 'script', 'style']:
            return

        # Пропускаємо елементи з класами навігації
        element_class = element.get('class')
        if element_class:
            classes = ' '.join(element_class).lower()
            skip_classes = ['sidebar', 'toctree', 'breadcrumb', 'navigation', 'menu',
                            'navbar', 'btn', 'button', 'prev-next', 'footer', 'header']
            if any(skip in classes for skip in skip_classes):
                return

        # Обробляємо текстові блоки
        if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Перевіряємо чи є в параграфі посилання-кнопки
            links = element.find_all('a')
            if links:
                # Якщо всі посилання мають клас кнопки - пропускаємо
                all_buttons = all(
                    link.get('class') and
                    any(btn in ' '.join(link.get('class')).lower()
                        for btn in ['btn', 'button'])
                    for link in links if link.get('class')
                )
                if all_buttons and len(links) > 0:
                    return

            text = element.get_text(strip=True)
            # Фільтруємо короткі тексти які можуть бути частиною навігації
            if text and len(text) > 2:
                text_list.append({
                    'type': element.name,
                    'text': text
                })

        # Обробляємо списки
        elif element.name == 'li':
            # Пропускаємо елементи списку в навігації
            parent_ul = element.find_parent(['ul', 'ol'])
            if parent_ul:
                parent_class = parent_ul.get('class')
                if parent_class:
                    classes = ' '.join(parent_class).lower()
                    if any(skip in classes for skip in ['toctree', 'menu', 'nav']):
                        return

            text = element.get_text(strip=True)
            if text and len(text) > 2:
                text_list.append({
                    'type': 'li',
                    'text': text
                })

        # Обробляємо блоки коду
        elif element.name in ['pre', 'code']:
            # Для code всередині pre - пропускаємо (обробимо pre)
            if element.name == 'code' and element.find_parent('pre'):
                return

            text = element.get_text(strip=True)
            if text and len(text) > 2:
                text_list.append({
                    'type': element.name,
                    'text': text
                })

        # Обробляємо blockquote
        elif element.name == 'blockquote':
            text = element.get_text(strip=True)
            if text and len(text) > 10:
                text_list.append({
                    'type': 'blockquote',
                    'text': text
                })

        # Рекурсивно обробляємо дочірні елементи для div, section тощо
        elif element.name in ['div', 'section', 'article', 'main', 'body']:
            for child in element.children:
                if hasattr(child, 'name'):
                    self._process_element(child, text_list, element.name)

    def add_to_document(self, content_data):
        """Додає контент до DOCX документу"""
        # Пропускаємо сторінки без контенту
        if not content_data['text']:
            return

        # Додаємо URL як коментар
        p = self.doc.add_paragraph()
        run = p.add_run(f"Джерело: {content_data['url']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Додаємо заголовок сторінки
        if content_data['title']:
            heading = self.doc.add_heading(content_data['title'], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Додаємо текстовий контент
        for item in content_data['text']:
            if item['type'] in ['h1', 'h2', 'h3', 'h4']:
                level = int(item['type'][1]) + 1  # h1 -> level 2, h2 -> level 3, etc.
                self.doc.add_heading(item['text'], level=min(level, 9))
            elif item['type'] in ['pre', 'code']:
                p = self.doc.add_paragraph(item['text'])
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(50, 50, 50)
            elif item['type'] == 'blockquote':
                p = self.doc.add_paragraph(item['text'])
                p.paragraph_format.left_indent = Pt(36)
                p.paragraph_format.right_indent = Pt(36)
                for run in p.runs:
                    run.font.italic = True
            elif item['type'] == 'li':
                self.doc.add_paragraph(item['text'], style='List Bullet')
            else:
                # Звичайний параграф
                if len(item['text']) > 5:  # Фільтруємо дуже короткі тексти
                    self.doc.add_paragraph(item['text'])

        # Додаємо роздільник
        self.doc.add_paragraph('_' * 80)
        self.doc.add_paragraph()

    def scrape(self):
        """Основний метод для скрапінгу"""
        print(f"Початок скрапінгу: {self.start_url}")
        print(f"Спочатку збираємо всі URLs в правильному порядку...\n")

        # Фаза 1: Збираємо всі URLs в правильному порядку
        self._collect_urls_recursive(self.start_url)

        total_pages = len(self.ordered_urls)
        print(f"\nЗнайдено {total_pages} сторінок")
        print(f"Починаємо завантаження контенту...\n")

        # Фаза 2: Завантажуємо контент в зібраному порядку
        for idx, url in enumerate(self.ordered_urls, 1):
            try:
                print(f"Обробка [{idx}/{total_pages}]: {url}")

                response = self.session.get(url, timeout=30)
                response.raise_for_status()

                soup = BeautifulSoup(response.content, 'html.parser')

                # Витягуємо контент
                content = self.extract_content(soup, url)
                self.add_to_document(content)

                # Затримка щоб не перевантажувати сервер
                time.sleep(1)

            except Exception as e:
                print(f"Помилка при обробці {url}: {str(e)}")
                continue

        print(f"\n✓ Завершено! Оброблено сторінок: {total_pages}")
        return total_pages

    def _collect_urls_recursive(self, url, depth=0):
        """Рекурсивно збирає URLs в порядку появи в навігації"""
        if url in self.visited or depth > 50:  # Обмеження глибини
            return

        self.visited.add(url)
        self.ordered_urls.append(url)

        if depth == 0:
            print(f"Сканування структури [{len(self.ordered_urls)}]: {url}")

        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')

            # Витягуємо посилання в порядку появи
            links = self.extract_links(soup, url)

            # Рекурсивно обходимо кожне посилання
            for link in links:
                if link not in self.visited:
                    if depth < 2:  # Показуємо прогрес тільки для перших рівнів
                        print(f"Сканування структури [{len(self.ordered_urls)}]: {link}")
                    self._collect_urls_recursive(link, depth + 1)

            time.sleep(0.5)  # Коротша затримка при скануванні структури

        except Exception as e:
            print(f"Помилка при скануванні {url}: {str(e)}")

    def save(self, filename='odoo_documentation.docx'):
        """Зберігає документ"""
        self.doc.save(filename)
        print(f"✓ Документ збережено: {filename}")


def main():
    start_url = "https://www.odoo.com/documentation/18.0/contributing.html"

    scraper = OdooDocsScraper(start_url)
    scraper.scrape()
    scraper.save('odoo_18_contributing_docs.docx')


if __name__ == "__main__":
    main()